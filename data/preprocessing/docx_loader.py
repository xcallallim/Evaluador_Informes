"""Carga especializada para documentos DOCX."""

from __future__ import annotations

from typing import List

from core.logger import log_error, log_info
from data.models.document import Document
from data.preprocessing.metadata import LoaderContext
from data.preprocessing.txt_loader import split_text_into_pages


def _flush_page(pages: List[str], current_page: List[str]) -> None:
    if current_page:
        pages.append("\n".join(current_page))
        current_page.clear()


def load(filepath: str, *, extract_tables: bool = True) -> Document:
    """Carga archivos ``.docx`` devolviendo un ``Document`` parcial."""

    log_info("📄 Cargando archivo DOCX...")

    try:
        import docx  # type: ignore

        doc = docx.Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n".join(paragraphs)

        pages: List[str] = []
        current_page: List[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                current_page.append(text)

            element = getattr(paragraph, "_p", None)
            if element is None:
                continue

            nsmap = getattr(element, "nsmap", {}) or {}
            w_namespace = nsmap.get("w")
            if not w_namespace:
                continue

            for br in element.findall(f".//{{{w_namespace}}}br"):
                if br.get(f"{{{w_namespace}}}type") == "page":
                    _flush_page(pages, current_page)
                    break

        _flush_page(pages, current_page)

        if not pages:
            pages = split_text_into_pages(raw_text, max_chars=4000)

        pages_with_tags = [f"=== PAGE {index + 1} ===\n{page}" for index, page in enumerate(pages)]
        full_text_with_pages = "\n\n".join(pages_with_tags)

        tables: List[List[str]] = []
        if extract_tables:
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):
                        table_data.append(row_cells)
                if table_data:
                    tables.append(table_data)

        context = LoaderContext(
            tables_meta={"docx": tables} if tables else {},
            extra_metadata={"is_ocr": False, "extraction_method": "docx"},
        )

        metadata = {"loader_context": context.as_dict()}

        return Document(
            content=full_text_with_pages,
            pages=pages,
            tables=tables if extract_tables else [],
            metadata=metadata,
        )
    except Exception as exc:  # pragma: no cover - log y re-raise
        log_error(f"❌ Error cargando DOCX: {exc}")
        raise