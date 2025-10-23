# data/preprocessing/loader.py
# ------------------------------------------------------------
# DocumentLoader: Carga TXT, DOCX, PDF (digital + OCR)
# - DOCX: extrae texto y tablas
# - PDF digital: extrae texto con pdfplumber
# - PDF escaneado: extrae texto con OCR (Tesseract + OpenCV)
# - PDF tablas: extrae tablas con Camelot (opcional) y exporta a CSV
# - Metadatos consistentes y listos para el pipeline
# ------------------------------------------------------------

import os
import re
import shutil
import uuid
import fitz
from typing import Any, Dict, List, Optional, Tuple

from core.logger import log_error, log_info, log_warn
from core.config import TESSERACT_PATH
from core.utils import clean_spaces, ensure_dir
from data.models.document import Document

# No hacemos "from core.config import *" para mantener claridad.
try:
    from core.config import GHOSTSCRIPT_PATH
except Exception:
    GHOSTSCRIPT_PATH = None


MEANINGFUL_TEXT_PATTERN = re.compile(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9]")


class DocumentLoader:
    def __init__(self):
        # Extensiones soportadas
        self.supported_extensions = [".pdf", ".docx", ".txt"]
        # Flag interno para saber si la última carga PDF usó OCR.
        self._last_pdf_was_ocr: bool = False
        # Texto crudo más reciente extraído de un PDF (digital u OCR).
        self._last_pdf_raw_text: str = ""
        # Lista de issues detectadas durante la carga actual.
        self._issues: List[str] = []
        # Configurar Tesseract si está disponible
        self._configure_tesseract()
        # Resolver ruta de Ghostscript si el entorno no permite instalación global.
        self.ghostscript_cmd = self._resolve_ghostscript()

    # ------------------------------------------------------------
    # Configuración OCR
    # ------------------------------------------------------------
    def _configure_tesseract(self):
        """Configura la ruta de Tesseract OCR si está disponible."""
        try:
            import pytesseract
            if TESSERACT_PATH and os.path.exists(TESSERACT_PATH):
                pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
                log_info(f"Tesseract configurado en: {TESSERACT_PATH}")
            else:
                log_warn("Ruta de Tesseract no encontrada o inválida. OCR deshabilitado.")
        except ImportError:
            log_warn("pytesseract no está instalado. OCR deshabilitado.")
    
    def _resolve_ghostscript(self) -> Optional[str]:
        """Localiza Ghostscript permitiendo binarios descargados sin privilegios sudo."""
        candidates = [
            os.environ.get("GHOSTSCRIPT_PATH"),
            os.environ.get("GS_BIN"),
            GHOSTSCRIPT_PATH,
        ]

        for executable in ("gs", "gswin64c", "gswin32c"):
            path = shutil.which(executable)
            if path:
                candidates.append(path)

        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                log_info(f"Ghostscript detectado en: {candidate}")
                return candidate

        log_warn(
            "Ghostscript no está disponible. Camelot seguirá intentando, "
            "pero la extracción de tablas puede fallar sin este binario."
        )
        return None
    
    def _register_issue(self, message: str) -> None:
        """Acumula advertencias de la carga actual sin duplicados."""

        if message and message not in self._issues:
            self._issues.append(message)

    # ------------------------------------------------------------
    # API principal
    # ------------------------------------------------------------
    def load(self, filepath: str, extract_tables: bool = True, extract_images: bool = True) -> Document:

        """
        Carga un archivo y devuelve un objeto Document.

        Args:
            filepath: ruta del archivo a cargar.
            extract_tables: si True, intenta extraer tablas (DOCX y PDF con Camelot).
                            Esta opción es ideal para habilitar/deshabilitar desde una UI.
        """
        # Reiniciar flag OCR para cada carga
        self._last_pdf_was_ocr = False
        self._last_pdf_raw_text = ""
        self._issues = []

        filepath = os.fspath(filepath)
        
        if not os.path.isfile(filepath):
            log_error(f"Archivo no encontrado: {filepath}")
            raise FileNotFoundError(f"Archivo no encontrado: {filepath}")

        _, ext = os.path.splitext(filepath)
        ext = ext.lower()

        if ext not in self.supported_extensions:
            log_error(f"Formato no soportado: {ext}")
            raise ValueError(f"Formato no soportado: {ext}")

        log_info(f"Cargando archivo: {filepath}")

        # Estructura de tablas en metadata (multi-origen)
        tables_meta: Dict[str, Any] = {}
        flat_tables: List[Any] = []

        if ext == ".txt":
            content, pages, txt_tables = self._load_txt(filepath)
            if extract_tables and txt_tables:
                tables_meta["txt"] = txt_tables
                flat_tables.extend(txt_tables)

        elif ext == ".docx":
            content, pages, docx_tables = self._load_docx(filepath)
            if extract_tables and docx_tables:
                tables_meta["docx"] = docx_tables
                flat_tables.extend(docx_tables)

        elif ext == ".pdf":
            content, pages = self._load_pdf(filepath)

            # Cuando la extracción usa OCR debemos mantener ``content`` vacío.
            # El texto reconocido se entrega únicamente en ``pages`` y metadata.
            if self._last_pdf_was_ocr:
                content = ""

            if extract_tables:
                pdf_tables_meta = self._extract_pdf_tables(filepath)
                if pdf_tables_meta:
                    tables_meta["pdf"] = pdf_tables_meta
                    flat_tables.extend(pdf_tables_meta)

            if extract_images:
                pdf_images_meta = self._extract_pdf_images(filepath)

        else:
            # Por seguridad; no debería alcanzarse por validación previa.
            raise NotImplementedError(f"Carga para {ext} aún no implementada.")

        metadata = {
            "filename": os.path.basename(filepath),
            "extension": ext,
            "source": filepath,
            "pages": pages,             # lista de textos por página (PDF) o None
            "tables": tables_meta or {},# {"docx": [...], "pdf": [{"page": int, "path": str}, ...]}
            "processed_with": "DocumentLoader",
        }

        if ext == ".pdf":
            metadata["is_ocr"] = self._last_pdf_was_ocr
            metadata["extraction_method"] = "ocr" if self._last_pdf_was_ocr else "embedded_text"
            metadata["raw_text"] = self._last_pdf_raw_text
        elif ext == ".docx":
            metadata["is_ocr"] = False
            metadata["extraction_method"] = "docx"
        else:
            metadata["is_ocr"] = False
            metadata["extraction_method"] = "text"

        # Añadir imágenes al metadata
        if extract_images:
            metadata["images"] = pdf_images_meta if 'pdf_images_meta' in locals() else []

        if self._issues:
            metadata["issues"] = list(self._issues)

        # --- LOG RESUMEN DE CARGA ---
        log_info("📊 Resumen de carga del documento:")
        pages_count = len(pages) if isinstance(pages, list) else 0
        log_info(f"   • Páginas detectadas: {pages_count}")
        log_info(f"   • Tablas extraídas: {len(tables_meta.get('pdf', [])) + len(tables_meta.get('docx', []))}")
        log_info(f"   • Imágenes extraídas: {len(metadata.get('images', [])) if extract_images else 0}")
        if self._issues:
            log_warn("⚠ Documento cargado con advertencias. Revisar detalles:")
            for issue in self._issues:
                log_warn(f"   • {issue}")
        else:
            log_info("✅ Documento cargado correctamente ✅")

        # Devolvemos Document con texto completo + metadata rica
        return Document(
            content=content,
            metadata=metadata,
            pages=pages,
            tables=flat_tables,
        )

    def _split_text_into_pages(self, text: str, max_chars: int = 4000) -> List[str]:
        """
        Divide texto largo en 'páginas' aproximadas si no hay saltos naturales.
        Solo como fallback para TXT simples.
        """
        return [text[i:i + max_chars] for i in range(0, len(text), max_chars)]
    
    def _slugify_filename(self, name: str) -> str:
        """
        Convierte un nombre a un 'slug' seguro para carpetas/archivos.
        Temporal aquí; luego podemos moverlo a core/utils.py.
        """
        # Normaliza, quita tildes y caracteres raros
        import unicodedata
        name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("ascii")
        name = name.lower()
        name = re.sub(r"[^a-z0-9]+", "_", name)
        return name.strip("_")


    # ------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------
    def _load_txt(self, filepath: str) -> Tuple[str, List[str], List]:
        """
        Carga archivos .txt conservando estructura
        Devuelve: texto completo, lista de páginas y tablas vacías (TXT no tiene tablas).
        """
        log_info("📄 Leyendo archivo TXT...")

        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            # Dividir contenido en páginas si hay señales de salto
            pages = re.split(r"\f|\n=== PAGE \d+ ===", content)
            if len(pages) <= 1:
                # Si no hay saltos claros, intenta dividir por longitud como fallback
                pages = self._split_text_into_pages(content, max_chars=4000)

            # Añadir separador de página visible
            pages_with_tags = [f"=== PAGE {i+1} ===\n{p.strip()}" for i, p in enumerate(pages)]
            full_text_with_pages = "\n\n".join(pages_with_tags)

            return full_text_with_pages, pages, []  # ← importante: [] para tablas

        except Exception as e:
            log_error(f"❌ Error leyendo archivo TXT: {e}")
            raise

    # ------------------------------------------------------------
    # DOCX (texto + tablas)
    # ------------------------------------------------------------
    def _load_docx(self, filepath: str) -> Tuple[str, List[str], List]:
        """
        Carga texto y tablas desde un archivo Word (.docx),
        respetando saltos de página visibles y extrayendo tablas.
        Devuelve:
        - texto completo
        - lista de páginas (si no encuentra marcas de página, usa fallback simple)
        - tablas en estructura FILAS x COLUMNAS
        """
        log_info("📄 Cargando archivo DOCX...")

        try:
            import docx
            doc = docx.Document(filepath)

            # Extraer texto respetando párrafos
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            raw_text = "\n".join(paragraphs)

            # Detectar saltos de página utilizando marcadores explícitos de Word
            pages = []
            current_page: List[str] = []

            def _flush_page() -> None:
                if current_page:
                    pages.append("\n".join(current_page))
                    current_page.clear()
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    current_page.append(text)

                element = getattr(paragraph, "_p", None)
                if element is None:
                    continue

                nsmap = getattr(element, "nsmap", {}) or {}
                w_namespace = nsmap.get("w")
                if not w_namespace:
                    continue

                # Buscar cualquier salto de página declarado en el XML del párrafo
                for br in element.findall(f".//{{{w_namespace}}}br"):
                    br_type = br.get(f"{{{w_namespace}}}type")
                    if br_type == "page":
                        _flush_page()
                        break

            _flush_page()

            # Fallback: si no detectó páginas correctamente
            if not pages:
                pages = self._split_text_into_pages(raw_text, max_chars=4000)

            # Insertar separadores visibles
            pages_with_tags = [f"=== PAGE {i+1} ===\n{p}" for i, p in enumerate(pages)]
            full_text_with_pages = "\n\n".join(pages_with_tags)

            # Extraer tablas
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):
                        table_data.append(row_cells)
                if table_data:
                    tables.append(table_data)

            log_info(f"📄 DOCX cargado correctamente: {len(pages)} páginas simuladas y {len(tables)} tablas detectadas.")
            return full_text_with_pages, pages, tables

        except Exception as e:
            log_error(f"❌ Error cargando DOCX: {e}")
            raise


    # ------------------------------------------------------------
    # PDF (detección digital vs escaneado + OCR preciso)
    # ------------------------------------------------------------
    def _load_pdf(self, filepath: str) -> Tuple[str, List[str]]:
        """
        Detecta si el PDF es digital (texto embebido) o escaneado (imágenes),
        y aplica la extracción adecuada.
        Devuelve:
        - full_text: texto completo con separadores de página
        - pages: lista con texto de cada página
        """
        log_info("📄 Analizando PDF...")

        # 1) Intentar primero con PyMuPDF: es rápido y fiable para detectar capa de texto.
        pymupdf_detection = self._try_load_pdf_with_pymupdf(filepath)
        pymupdf_detected: Optional[bool] = None
        pymupdf_pages: List[str] = []

        if pymupdf_detection:
            pymupdf_detected, pymupdf_pages = pymupdf_detection

        if pymupdf_detected is False:
            log_warn("⚠ PDF escaneado detectado (PyMuPDF). Ejecutando OCR...")
            self._last_pdf_was_ocr = True
            ocr_full_text, pages_text = self._load_pdf_ocr(filepath)
            self._last_pdf_raw_text = ocr_full_text
            return "", pages_text

        pages_text: List[str] = []
        is_digital = False
        pdfplumber_failed = False

        # 2) Intentar PDF digital con pdfplumber
        try:
            import pdfplumber  # type: ignore
        except Exception:
            log_warn(
                "pdfplumber no está disponible o falló la importación. "
                "Se procederá al OCR si no se detecta texto con otros métodos."
            )
        else:
            if pymupdf_detected is not False:
                try:
                    with pdfplumber.open(filepath) as pdf:
                        for page in pdf.pages:
                            text = page.extract_text() or ""
                            words = []
                            try:
                                words = page.extract_words() or []
                            except Exception:
                                # Algunos backends pueden fallar al extraer palabras; continuar silenciosamente.
                                words = []

                            if text.strip() or words:
                                is_digital = True

                            if not text.strip() and words:
                                # Fallback para PDFs digitales donde extract_text falla pero sí hay palabras.
                                text = " ".join(word.get("text", "") for word in words).strip()

                            if not text.strip() and not words:
                                # Algunos motores devuelven caracteres individuales accesibles via page.chars
                                # aunque extract_text y extract_words queden vacíos. Reconstruimos líneas simples.
                                chars = getattr(page, "chars", []) or []
                                if chars:
                                    is_digital = True

                                    # Ordenar por coordenadas para agrupar líneas aproximadas.
                                    sorted_chars = sorted(
                                        chars,
                                        key=lambda c: (round(c.get("top", 0.0), 1), c.get("x0", 0.0)),
                                    )

                                    from itertools import groupby

                                    lines = []
                                    for _, group in groupby(
                                        sorted_chars, key=lambda c: round(c.get("top", 0.0), 1)
                                    ):
                                        line_text = "".join(char.get("text", "") for char in group)
                                        if line_text.strip():
                                            lines.append(line_text)

                                    if lines:
                                        text = "\n".join(lines)

                            # Respetar saltos tal cual vienen (limpieza ocurrirá en Cleaner)
                            pages_text.append(text)
                except Exception as e:
                    log_warn(
                        "No se pudo leer como PDF digital con pdfplumber. "
                        f"Detalle: {e}. Se procederá al OCR si no hay más opciones."
                    )
                    pages_text = []
                    pdfplumber_failed = True

        # 3) PDF DIGITAL → devolver
        if is_digital:
            log_info("✅ PDF digital detectado (pdfplumber)")
            pages_with_tags = [f"=== PAGE {i+1} ===\n{p.strip()}" for i, p in enumerate(pages_text)]
            full_text = "\n\n".join(pages_with_tags)
            self._last_pdf_was_ocr = False
            self._last_pdf_raw_text = full_text
            return full_text, pages_text
        
        if pymupdf_detected:
            meaningful_pages = [page for page in pymupdf_pages if page.strip()]
            if meaningful_pages:
                log_info("✅ PDF digital detectado (PyMuPDF fallback)")
                pages_with_tags = [
                    f"=== PAGE {i+1} ===\n{p.strip()}" for i, p in enumerate(pymupdf_pages)
                ]
                full_text = "\n\n".join(pages_with_tags)
                self._last_pdf_was_ocr = False
                self._last_pdf_raw_text = full_text
                return full_text, pymupdf_pages

        # 4) Sino → usar OCR
        if pdfplumber_failed:
            log_warn(
                "⚠ PDF escaneado detectado. Ejecutando OCR tras fallo de pdfplumber..."
            )
        else:
            log_warn("⚠ PDF escaneado detectado. Ejecutando OCR...")
        self._last_pdf_was_ocr = True
        ocr_full_text, pages_text = self._load_pdf_ocr(filepath)
        self._last_pdf_raw_text = ocr_full_text
        # Para el pipeline oficial (loader → cleaner → segmenter → splitter),
        # mantenemos el contenido vacío y dejamos el texto crudo en metadata/pages.
        return "", pages_text

    def _try_load_pdf_with_pymupdf(self, filepath: str) -> Optional[Tuple[bool, List[str]]]:
        """Analiza el PDF con PyMuPDF para detectar si existe capa de texto real."""

        try:
            import fitz  # type: ignore
        except Exception:
            log_warn(
                "PyMuPDF no está disponible para el fallback de PDFs digitales."
            )
            return None

        try:
            pages_text: List[str] = []
            pages_with_content = 0

            with fitz.open(filepath) as pdf_document:
                for page in pdf_document:
                    page_text = self._extract_text_from_pymupdf_page(page)
                    if page_text.strip():
                        pages_with_content += 1
                    pages_text.append(page_text)

            if pages_with_content == 0:
                return (False, pages_text)

            return (True, pages_text)

        except Exception as e:
            log_warn(
                "Fallback con PyMuPDF falló. Se procederá a OCR. "
                f"Detalle: {e}"
            )
            return None


    def _extract_text_from_pymupdf_page(self, page: "fitz.Page") -> str:
        """Obtiene texto significativo de una página PyMuPDF evitando falsos positivos."""

        try:
            dict_data = page.get_text("dict") or {}
        except Exception:
            dict_data = {}

        lines: List[str] = []
        page_has_meaningful = False

        blocks = dict_data.get("blocks", []) if isinstance(dict_data, dict) else []
        for block in blocks:
            if not isinstance(block, dict) or block.get("type") != 0:
                continue

            for line in block.get("lines", []) or []:
                if not isinstance(line, dict):
                    continue
                spans_text: List[str] = []
                for span in line.get("spans", []) or []:
                    if not isinstance(span, dict):
                        continue
                    span_text = span.get("text", "")
                    if not span_text:
                        continue
                    cleaned = span_text.replace("\u00ad", "").strip()
                    if not cleaned:
                        continue
                    spans_text.append(cleaned)
                    if not page_has_meaningful and MEANINGFUL_TEXT_PATTERN.search(cleaned):
                        page_has_meaningful = True

                if spans_text:
                    lines.append(" ".join(spans_text))

        text = "\n".join(lines).strip()

        if not text:
            # Fallback a extracción básica para cubrir casos donde no hay estructura dict
            try:
                fallback = page.get_text("text") or ""
                fallback = fallback.strip()
                if fallback and MEANINGFUL_TEXT_PATTERN.search(fallback):
                    text = fallback
                    page_has_meaningful = True
            except Exception:
                text = ""

        if not text or not page_has_meaningful:
            return ""

        return text


    def _load_pdf_ocr(self, filepath: str) -> Tuple[str, List[str]]:
        """
        OCR de PDF escaneado con PyMuPDF + OpenCV + Tesseract.
        Perfil 'limpio' (rápido, conserva estructura). Manejo básico de columnas con psm=4.
        """
        try:
            import fitz  # PyMuPDF
        except Exception as import_error:
            log_error(f"PyMuPDF es requerido para OCR: {import_error}")
            self._register_issue("PyMuPDF no está disponible para ejecutar OCR.")
            return "", []
        import cv2
        import numpy as np
        import pytesseract
        from PIL import Image
        import tempfile

        def _cleanup_tmp(paths: List[str]) -> None:
            for path in paths:
                if not path:
                    continue
                try:
                    os.remove(path)
                except FileNotFoundError:
                    continue
                except OSError:
                    # En Windows puede haber procesos que mantengan el handler un instante.
                    pass

        pages_text: List[str] = []
        log_info("🧠 Iniciando OCR página por página...")

        previous_tempdir = getattr(pytesseract.pytesseract, "TEMP_DIR", None)

        try:
            with tempfile.TemporaryDirectory(prefix="tess_tmp_") as tess_tempdir:
                pytesseract.pytesseract.TEMP_DIR = tess_tempdir

                pdf_document = fitz.open(filepath)
                total = len(pdf_document)

                for i, page in enumerate(pdf_document):
                    if i % 5 == 0:
                        log_info(f"OCR procesando página {i+1}/{total}")

                    page_text = ""
                    image_path = ""
                    text_path = ""
                    pil_image: Optional[Image.Image] = None
                    try:
                        pix = page.get_pixmap(dpi=200)
                        img = np.frombuffer(
                            pix.samples, dtype=np.uint8
                        ).reshape(pix.height, pix.width, pix.n)

                        if pix.n == 4:
                            img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)
                        elif pix.n == 1:
                            img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)

                        del pix

                        # Preproceso básico para OCR
                        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                        gray = cv2.bilateralFilter(gray, 11, 17, 17)
                        thresh = cv2.adaptiveThreshold(
                            gray,
                            255,
                            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                            cv2.THRESH_BINARY,
                            11,
                            2,
                        )

                        pil_image = Image.fromarray(thresh)

                        try:
                            # Guardamos la imagen en disco para evitar locks de Windows con NamedTemporaryFile.
                            base_name = os.path.join(
                                tess_tempdir,
                                f"page_{i+1}_{uuid.uuid4().hex}",
                            )
                            image_path = f"{base_name}.png"
                            text_path = f"{base_name}.txt"
                            pil_image.save(image_path)

                            pytesseract.pytesseract.run_tesseract(
                                image_path,
                                base_name,
                                extension="txt",
                                lang="spa",
                                config="--psm 4",
                            )

                            if os.path.exists(text_path):
                                with open(text_path, "r", encoding="utf-8", errors="ignore") as fh:
                                    page_text = fh.read()
                            else:
                                self._register_issue(
                                    f"El OCR no devolvió texto en la página {i+1}."
                                )
                        except PermissionError as page_error:
                            log_warn(f"OCR falló en la página {i+1}: {page_error}")
                            self._register_issue(
                                f"OCR falló en la página {i+1}."
                            )
                        except Exception as page_error:
                            log_warn(f"OCR falló en la página {i+1}: {page_error}")
                            self._register_issue(
                                f"OCR falló en la página {i+1}."
                            )
                    finally:
                        if pil_image is not None:
                            pil_image.close()
                        _cleanup_tmp([image_path, text_path])

                    pages_text.append(page_text)

                pdf_document.close()

            # Concatenar con separadores de página
            pages_with_tags = [f"=== PAGE {i+1} ===\n{p}" for i, p in enumerate(pages_text)]
            full_text = "\n\n".join(pages_with_tags)

            log_info("✅ OCR completado correctamente.")
            return full_text, pages_text

        except Exception as e:
            log_error(f"❌ Error durante OCR: {e}")
            self._register_issue("El OCR no pudo procesar el PDF escaneado.")

            fallback_pages: List[str] = []
            try:
                with fitz.open(filepath) as pdf_document:
                    fallback_pages = [""] * len(pdf_document)
            except Exception:
                fallback_pages = []

            return "", fallback_pages
        finally:
            if previous_tempdir is not None:
                pytesseract.pytesseract.TEMP_DIR = previous_tempdir
            elif hasattr(pytesseract.pytesseract, "TEMP_DIR"):
                delattr(pytesseract.pytesseract, "TEMP_DIR")



    # ------------------------------------------------------------
    # Tablas PDF con Camelot (export a CSV por documento)
    # ------------------------------------------------------------
    def _save_camelot_tables(
        self,
        tables: Any,
        base_dir: str,
        tables_meta: List[Dict[str, Any]],
    ) -> int:
        """Guarda tablas de Camelot en CSV y actualiza metadata.

        Devuelve el número de tablas válidas persistidas.
        """

        if not tables:
            return 0

        valid_count = 0

        for idx, table in enumerate(tables):
            try:
                df = getattr(table, "df", None)
                if df is None or df.empty:
                    continue

                if df.shape[0] < 2 or df.shape[1] < 2:
                    # Filtrar tablas triviales o con datos incompletos.
                    continue

                page_number = getattr(table, "page", None)
                if page_number is None:
                    parsing_report = getattr(table, "parsing_report", None)
                    if isinstance(parsing_report, dict):
                        page_number = parsing_report.get("page")

                if not page_number:
                    page_number = "unknown"

                out_name = f"page_{page_number}_{idx + 1}.csv"
                out_path = os.path.join(base_dir, out_name)
                ensure_dir(base_dir)
                df.to_csv(out_path, index=False)

                tables_meta.append(
                    {
                        "page": page_number,
                        "path": out_path.replace("\\", "/"),
                    }
                )
                valid_count += 1
            except Exception as exc:  # pragma: no cover - protección adicional
                log_warn(f"No se pudo guardar tabla {idx + 1}: {exc}")
                self._register_issue(
                    f"Fallo al exportar tabla {idx + 1} detectada por Camelot"
                )

        return valid_count
    
    
    def _extract_pdf_tables(self, filepath: str) -> List[Dict[str, Any]]:
        """
        Extrae tablas de un PDF usando Camelot.
        - Usa Ghostscript si está configurado (GHOSTSCRIPT_PATH).
        - Intenta primero 'lattice' (tablas con líneas) y luego 'stream' (tablas sin líneas).
        - Filtra tablas válidas (>= 2 filas y >= 2 columnas).
        - Exporta cada tabla a CSV en: data/outputs/tables/<nombre_doc>/page_<n>_<idx>.csv
        - Devuelve lista de dicts: [{"page": int, "path": str}, ...]
        """
        tables_meta: List[Dict[str, Any]] = []

        # 0) Validar Camelot instalado
        try:
            import camelot  # type: ignore
        except Exception:
            log_warn("Camelot no está instalado. Saltando extracción de tablas PDF.")
            self._register_issue("Camelot no está disponible; no se extrajeron tablas del PDF.")
            return tables_meta

        # 1) Construir carpeta de salida: data/outputs/tables/<nombre_sin_ext>/
        fname = os.path.basename(filepath)
        stem, _ = os.path.splitext(fname)
        base_dir = os.path.join("data", "outputs", "tables", stem)
        ensure_dir(base_dir)

        # 2) Intentar extracción con 'lattice' (mejor para tablas con bordes marcados)
        try:
            log_info("Extrayendo tablas (Camelot - lattice)...")
            kwargs = {"flavor": "lattice", "pages": "all"}
            if self.ghostscript_cmd:
                kwargs["gs"] = self.ghostscript_cmd

            tables = camelot.read_pdf(filepath, **kwargs)
            valid_count = self._save_camelot_tables(tables, base_dir, tables_meta)
        except Exception as e:
            log_warn(f"Fallo 'lattice': {e}")
            self._register_issue("Camelot falló con el modo lattice.")
            valid_count = 0

        # 3) Si no encontró nada útil, intentamos 'stream' (detecta columnas por espacios)
        if valid_count == 0:
            try:
                log_info("Reintentando tablas (Camelot - stream)...")
                kwargs = {"flavor": "stream", "pages": "all"}
                if self.ghostscript_cmd:
                    kwargs["gs"] = self.ghostscript_cmd

                tables = camelot.read_pdf(filepath, **kwargs)
                valid_count = self._save_camelot_tables(tables, base_dir, tables_meta)
            except Exception as e:
                log_warn(f"Fallo 'stream': {e}")
                self._register_issue("Camelot falló con el modo stream.")

        log_info(f"Tablas extraídas: {len(tables_meta)}")
        return tables_meta
    
    # ------------------
    # Imágenes PDF 
    # ------------------

    def _extract_pdf_images(self, filepath: str) -> List[Dict[str, Any]]:
        """
        Extrae imágenes grandes desde un PDF y las guarda en:
        data/outputs/images/<nombre_pdf_limpio>/page_<n>_img_<k>.png

        Reglas:
        - Solo guarda imágenes "útiles" (área > 30,000 px)
        - Usa nombres/carpetas 'seguros' (sin tildes ni símbolos raros)
        - Devuelve metadata: [{"page": int, "path": str}, ...]
        """
        images_meta: List[Dict[str, Any]] = []

        try:
            import fitz  # PyMuPDF
            from PIL import Image
            from io import BytesIO

            # 1) Carpeta base por PDF con nombre seguro
            fname = os.path.basename(filepath)
            stem, _ = os.path.splitext(fname)
            safe_stem = self._slugify_filename(stem)  # ⬅︎ util local (abajo)
            base_dir = os.path.join("data", "outputs", "images", safe_stem)
            ensure_dir(base_dir)

            # 2) Abrir PDF y recorrer páginas
            doc = fitz.open(filepath)
            for page_index in range(len(doc)):
                page = doc[page_index]
                img_list = page.get_images(full=True)  # [(xref, smth...), ...]

                img_counter = 0
                for img_info in img_list:
                    xref = img_info[0]
                    try:
                        pix = fitz.Pixmap(doc, xref)
                        # Convertir a RGB si es CMYK/GRAYA/etc
                        if pix.n > 4:  # CMYK u otros
                            pix = fitz.Pixmap(fitz.csRGB, pix)

                        # Filtro por tamaño (área > 30k px)
                        if (pix.width * pix.height) < 30000:
                            continue

                        img_counter += 1
                        img_bytes = pix.tobytes("png")
                        img = Image.open(BytesIO(img_bytes))

                        out_name = f"page_{page_index+1}_img_{img_counter}.png"
                        out_path = os.path.join(base_dir, out_name)
                        img.save(out_path, format="PNG")

                        images_meta.append({
                            "page": page_index + 1,
                            "path": out_path.replace("\\", "/")
                        })

                    except Exception as e:
                        log_warn(f"No se pudo extraer imagen (pág. {page_index+1}): {e}")
                        continue

            doc.close()

            log_info(f"🖼️ Imágenes extraídas: {len(images_meta)}")
            return images_meta

        except Exception as e:
            log_warn(f"No se pudieron extraer imágenes del PDF: {e}")
            return images_meta

